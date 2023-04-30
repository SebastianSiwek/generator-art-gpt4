# Importowanie potrzebnych modułów
import openai
from docx import Document
import re
import os

# Uzyskiwanie dostępu do modelu GPT-4
openai.api_key = "WPROWADŹ-SWÓJ-KLUCZ-API"

# Definiowanie zmiennych na podstawie inputów
industry = input("Wprowadź nazwę branży: ")
main_keyword = input("Wprowadź główne słowo kluczowe: ")
related_keywords = input("Wprowadź dodatkowe słowa kluczowe: ")

# Definiowanie funkcji odpowiedzialnej za generowanie tematu wpisu blogowego
def generate_topic():

    # Generowanie tematu blogowego
    generated_topic = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem w dziedzinie: {industry}"},
            {"role": "user", "content": f"Stwórz ciekawy, angażujący tytuł artykułu na podstawie podanego słowa kluczowego: {main_keyword}. Pisz w języku polskim. Nie umieszczaj cyfr w tytule. Nie umieszczaj tytułu w cudzyłowie."}
        ],
        max_tokens=200,
        temperature=0.8,
    )
    return generated_topic.choices[0].message.content

# Definiowanie funkcji odpowiedzialnej za generowanie spisu treści
def generate_outline(generated_topic):
    generated_outline = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem w dziedzinie: {industry}"},
            {"role": "user", "content": f"Zaplanuj i stwórz spis treści składający się z maksymalnie 5 punktów (bez podpunktów) dla artykułu na temat: {generated_topic}, pomiń wstęp i zakończenie. Pisz w języku polskim."}
        ],
        max_tokens=300,
        temperature=0.4,
    )
    return generated_outline.choices[0].message.content

# Definiowanie funkcji odpowiedzialnej za generowanie wstępu
def generate_intro(generated_topic):
    generated_intro = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem w dziedzinie: {industry}"},
            {"role": "user", "content": f"Napisz krótki, angażujący wstęp dla artykułu na temat: {generated_topic}. Pisz w języku polskim. Wygeneruj maksymalnie 5 zdań."}
        ],
        max_tokens=300,
        temperature=0.5,
    )
    return generated_intro.choices[0].message.content

# Definiowanie funkcji odpowiedzialnej za generowanie akapitów treści
def generate_paragraph(outline_points):
    paragraphs = []
    for point in outline_points:
        # Generowanie akapitów treści
        generated_paragraph = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Jesteś ekspertem w dziedzinie: {industry}"},
                {"role": "user", "content": f"Napisz akapit treści na temat: {point}. Wykorzystaj wskazane słowa kluczowe w tekście w formie nieodmienionej: {related_keywords}. Pisz w języku polskim. Wygeneruj maksymalnie 5 zdań."}
            ],
            max_tokens=300,
            temperature=0.5,
        )
        paragraphs.append(generated_paragraph.choices[0].message.content)
    return paragraphs

# Definiowanie funkcji odpowiedzialnej za generowanie podsumowania
def generate_summary(generated_topic):
    generated_summary = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem w dziedzinie: {industry}"},
            {"role": "user", "content": f"Napisz krótkie zakończenie dla artykułu na temat: {generated_topic}. Zachęć użytkownika do kreatywności w urządzaniu swojego wnętrza. Pisz w języku polskim. Wygeneruj maksymalnie 5 zdań."}
        ],
        max_tokens=300,
        temperature=0.5,
    )
    return generated_summary.choices[0].message.content

def create_blog_post_docx(generated_topic, generated_intro, paragraphs, generated_summary):
    # Tworzenie nowego dokumentu
    doc = Document()

    # Dodawanie tytułu wpisu jako nagłówka pierwszego poziomu
    doc.add_heading(generated_topic, level=1)

    # Dodawanie wstępu jako zwykłego paragrafu
    doc.add_paragraph(generated_intro)

    # Dodawanie na zmianę punktów konspektu jako nagłówka drugiego stopnia i odpowiadającego mu akapitu treści
    for i in range(len(outline_points)):
        doc.add_heading(outline_points[i], level=2)
        doc.add_paragraph(paragraphs[i])

    # Dodawanie nagłówka drugiego stopnia "Podsumowanie"
    doc.add_heading("Podsumowanie", level=2)
    doc.add_paragraph(generated_summary)

    # Usuwanie znaków specjalnych z wygenerowanego tematu
    def remove_special_chars(text):
        return re.sub(r'[^\w\s]', '', text)

    # Użycie wygenerowanego tytułu jako nazwy pliku (bez znaków specjalnych)
    file_name = remove_special_chars(generated_topic)

    # Ustalanie katalogu roboczego na katalog, w którym znajduje się skrypt
    script_dir = os.path.dirname(os.path.realpath(__file__))
    os.chdir(script_dir)

    folder_name = "Wygenerowane wpisy"

    # Sprawdzanie, czy folder istnieje; jeśli nie, tworzenie go
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    # Zapisywanie dokumentu
    file_path = os.path.join(folder_name, file_name + ".docx")
    doc.save(file_path)

# Główna część kodu
if __name__ == '__main__':
    # Generowanie tematu wpisu
    generated_topic = generate_topic()

    # Generowanie spisu treści
    generated_outline = generate_outline(generated_topic)

    # Generowanie wstępu
    generated_intro = generate_intro(generated_topic)

    # Generowanie spisu treści w postaci listy
    outline_points = re.findall(r'\d+\.\s(.*)', generated_outline)

    # Generowanie akapitów treści
    paragraphs = generate_paragraph(outline_points)

    # Generowanie podsumowania
    generated_summary = generate_summary(generated_topic)

    # Tworzenie dokumentu docx
    create_blog_post_docx(generated_topic, generated_intro, paragraphs, generated_summary)

    print("Dokument został utworzony.")